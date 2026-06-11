"""Build a personalized .docx report from report_payload.json.

HOST-SIDE (needs python-docx). Mirrors the structure of the SAMPLE template
(SAMPLE 00.00.2000.docx): cover + key-number table + birth chart + per-index
sections. Stored HTML content (<p>/<h4>/<ul><li>) is rendered to Word.

Usage:  python -m scripts.build_report_docx
Output: plans/reports/bao-cao-than-so-hoc-<name>-<dob>.docx
"""

from __future__ import annotations

import json
import re
import unicodedata
from html.parser import HTMLParser
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent
PAYLOAD = ROOT / "data" / "report_payload.json"
OUT_DIR = ROOT.parent / "plans" / "reports"


class _HtmlToDocx(HTMLParser):
    """Render simple <p>/<h4>/<ul><li> HTML into a python-docx document."""

    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.tag = None
        self.buf = ""

    def handle_starttag(self, tag, attrs):
        if tag in ("p", "h4", "li"):
            self.tag = tag
            self.buf = ""

    def handle_data(self, data):
        if self.tag:
            self.buf += data

    def handle_endtag(self, tag):
        if tag != self.tag:
            return
        text = re.sub(r"\s+", " ", self.buf).strip()
        if text:
            if tag == "h4":
                p = self.doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
            elif tag == "li":
                self.doc.add_paragraph(text, style="List Bullet")
            else:
                self.doc.add_paragraph(text)
        self.tag = None
        self.buf = ""


def render_html(doc, html: str):
    _HtmlToDocx(doc).feed(html or "")


def slug(s: str) -> str:
    s = s.replace("đ", "d").replace("Đ", "D")
    s = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode()
    return re.sub(r"[^a-z0-9]+", "-", s.lower()).strip("-")


def build():
    data = json.loads(PAYLOAD.read_text(encoding="utf-8"))
    header = data["header"]
    name = header["name"]
    dob_text = header["birth_day_text"]
    doc = docx.Document()

    # ── Cover ────────────────────────────────────────────────────────────
    h = doc.add_heading("THẦN SỐ HỌC PYTAGO", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("BÁO CÁO KHÁM PHÁ BẢN THÂN THÔNG QUA NHỮNG CON SỐ")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_table(rows=3, cols=2)
    info.style = "Light Grid Accent 1"
    for r, (k, v) in zip(info.rows, [("HỌ TÊN", name), ("NGÀY SINH", dob_text), ("ĐIỆN THOẠI", header.get("phone", ""))]):
        r.cells[0].text = k
        r.cells[1].text = str(v)

    # ── Key-number summary table ─────────────────────────────────────────
    doc.add_heading("TỔNG QUAN CÁC CHỈ SỐ", level=1)
    t = doc.add_table(rows=len(data["summary"]), cols=2)
    t.style = "Light List Accent 1"
    for row, (k, v) in zip(t.rows, data["summary"]):
        row.cells[0].text = str(k)
        row.cells[1].text = str(v)

    # ── Birth chart (Pythagoras 3x3) ─────────────────────────────────────
    doc.add_heading(f"BIỂU ĐỒ NGÀY SINH CỦA {name}", level=1)
    chart_tbl = doc.add_table(rows=3, cols=3)
    chart_tbl.style = "Table Grid"
    for ri, gr in enumerate(data["chart_grid"]):
        for ci, cell_text in enumerate(gr):
            cell = chart_tbl.rows[ri].cells[ci]
            cell.text = cell_text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for c in data.get("chart_sections", []):
        doc.add_heading(c["heading"], level=2)
        render_html(doc, c["content"])

    # ── Per-index sections ───────────────────────────────────────────────
    for s in data["sections"]:
        doc.add_heading(s["heading"], level=1)
        if s.get("title"):
            tp = doc.add_paragraph()
            run = tp.add_run(s["title"])
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        render_html(doc, s["content"])

    doc.add_paragraph()
    foot = doc.add_paragraph(
        "Nội dung bản Thần số học Pytago này chỉ mang giá trị tham khảo, "
        "khám phá bản thân."
    )
    foot.runs[0].italic = True

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / f"bao-cao-than-so-hoc-{slug(name)}-{re.sub(r'[^0-9]', '', dob_text)}.docx"
    doc.save(str(out))
    print(f"Wrote {out}  ({len(data['sections'])} sections)")


if __name__ == "__main__":
    import sys
    sys.stdout.reconfigure(encoding="utf-8")  # type: ignore[attr-defined]
    build()
