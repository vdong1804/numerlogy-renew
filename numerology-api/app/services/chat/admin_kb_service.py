"""Admin KB uploads: extract text from PDF/DOCX/TXT and ingest as a KB document.

Routes file uploads from the admin UI through the shared ``KbIngestionService``
so the embedding/chunk pipeline stays identical to the numerology backfill path.

source_type = "admin_upload"; source_ref = ``f"{admin_id}-{filename}"`` so an
admin re-uploading the same filename updates the existing document (idempotent).
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass
from typing import Optional

from sqlalchemy.ext.asyncio import AsyncSession

from app.db.models.chat.kb_document import KbDocument
from app.services.chat.kb_ingestion_service import KbIngestionService
from app.services.chat.pdf_parser_service import PdfParseError, PdfParserService

logger = logging.getLogger(__name__)

_PDF_EXTS = (".pdf",)
_DOCX_EXTS = (".docx",)
_TXT_EXTS = (".txt", ".md")
_SUPPORTED_EXTS = _PDF_EXTS + _DOCX_EXTS + _TXT_EXTS

# Max characters extracted per upload — guard against bombs / accidental huge files.
MAX_EXTRACTED_CHARS = 1_000_000


class UnsupportedFileType(ValueError):
    """File extension is not in the supported set."""


class ExtractedEmpty(ValueError):
    """File parsed without raising but produced no extractable text."""


@dataclass
class ExtractedDoc:
    text: str
    char_count: int
    file_kind: str  # "pdf" | "docx" | "txt"


class AdminKbService:
    """Admin-side wrapper around KbIngestionService for file uploads."""

    def __init__(
        self,
        db: AsyncSession,
        kb_ingestion: KbIngestionService,
        pdf_parser: Optional[PdfParserService] = None,
    ) -> None:
        self._db = db
        self._kb = kb_ingestion
        self._pdf = pdf_parser or PdfParserService()

    # -- Public --------------------------------------------------------------

    async def ingest_upload(
        self,
        filename: str,
        file_bytes: bytes,
        admin_id: int,
        title: Optional[str] = None,
    ) -> KbDocument:
        """Extract text + push through the KB ingestion pipeline."""
        extracted = self.extract(filename, file_bytes)
        if not extracted.text.strip():
            raise ExtractedEmpty("no extractable text in upload")

        source_ref = f"{admin_id}-{filename}"
        doc_title = title or filename
        metadata = {
            "filename": filename,
            "file_kind": extracted.file_kind,
            "char_count": extracted.char_count,
            "uploaded_by": admin_id,
        }
        doc = await self._kb.upsert_document(
            source_type="admin_upload",
            source_ref=source_ref,
            title=doc_title,
            content=extracted.text,
            metadata=metadata,
        )
        if admin_id is not None and doc.created_by is None:
            doc.created_by = admin_id
            await self._db.flush()
        logger.info(
            "admin_kb ingest: doc_id=%s ref=%s kind=%s chars=%d",
            doc.id, source_ref, extracted.file_kind, extracted.char_count,
        )
        return doc

    # -- Extraction ----------------------------------------------------------

    def extract(self, filename: str, file_bytes: bytes) -> ExtractedDoc:
        ext = _extension(filename)
        if ext not in _SUPPORTED_EXTS:
            raise UnsupportedFileType(
                f"unsupported file type {ext!r}; allowed: {_SUPPORTED_EXTS}"
            )
        if not file_bytes:
            raise ExtractedEmpty("file is empty")

        if ext in _PDF_EXTS:
            text = self._extract_pdf(file_bytes)
            kind = "pdf"
        elif ext in _DOCX_EXTS:
            text = self._extract_docx(file_bytes)
            kind = "docx"
        else:
            text = self._extract_txt(file_bytes)
            kind = "txt"

        if len(text) > MAX_EXTRACTED_CHARS:
            logger.warning(
                "admin_kb extract: truncating %s from %d to %d chars",
                filename, len(text), MAX_EXTRACTED_CHARS,
            )
            text = text[:MAX_EXTRACTED_CHARS]
        return ExtractedDoc(text=text, char_count=len(text), file_kind=kind)

    def _extract_pdf(self, data: bytes) -> str:
        try:
            pages = self._pdf.extract_pages(data)
        except PdfParseError as exc:
            raise ExtractedEmpty(str(exc)) from exc
        return PdfParserService.join_pages(pages)

    def _extract_docx(self, data: bytes) -> str:
        # python-docx import is lazy so the rest of the service stays usable
        # in environments where the lib isn't installed (e.g. PDF-only tests).
        try:
            import docx  # type: ignore[import-not-found]
        except ImportError as exc:  # pragma: no cover - guarded at startup
            raise RuntimeError(
                "python-docx not installed; add it to requirements"
            ) from exc

        try:
            document = docx.Document(io.BytesIO(data))
        except Exception as exc:  # noqa: BLE001
            raise ExtractedEmpty(f"could not open DOCX: {exc}") from exc

        parts: list[str] = []
        for para in document.paragraphs:
            text = (para.text or "").strip()
            if text:
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(
                    (cell.text or "").strip() for cell in row.cells if cell.text
                )
                if row_text:
                    parts.append(row_text)
        return "\n\n".join(parts)

    def _extract_txt(self, data: bytes) -> str:
        # Try UTF-8 first, fall back to latin-1 (lossless), then strip BOM.
        for encoding in ("utf-8-sig", "utf-8", "latin-1"):
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="replace")


def _extension(filename: str) -> str:
    idx = filename.rfind(".")
    return filename[idx:].lower() if idx >= 0 else ""


__all__ = [
    "AdminKbService",
    "ExtractedDoc",
    "ExtractedEmpty",
    "UnsupportedFileType",
    "MAX_EXTRACTED_CHARS",
]
